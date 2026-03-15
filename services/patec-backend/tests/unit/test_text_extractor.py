from pathlib import Path

import fitz
from docx import Document as DocxDocument
from openpyxl import Workbook

from app.services.text_extractor import extract_text


def test_extract_docx(tmp_path: Path):
    path = tmp_path / "sample.docx"
    doc = DocxDocument()
    doc.add_paragraph("Linha de teste DOCX")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    doc.save(path)

    extracted = extract_text(str(path), "docx")
    assert "Linha de teste DOCX" in extracted
    assert "A | B" in extracted


def test_extract_xlsx(tmp_path: Path):
    path = tmp_path / "sample.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "Planilha"
    ws.append(["coluna_1", "coluna_2"])
    ws.append(["v1", "v2"])
    wb.save(path)
    wb.close()

    extracted = extract_text(str(path), "xlsx")
    assert "Aba: Planilha" in extracted
    assert "coluna_1 | coluna_2" in extracted
    assert "v1 | v2" in extracted


def test_extract_pdf(tmp_path: Path):
    path = tmp_path / "sample.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Texto PDF para extracao")
    pdf.save(path)
    pdf.close()

    extracted = extract_text(str(path), "pdf")
    assert "Texto PDF para extracao" in extracted
