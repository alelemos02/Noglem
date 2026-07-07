import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import load_workbook


def _format_table(table_data: list[list]) -> str:
    """Format extracted table data as pipe-separated text with headers."""
    if not table_data:
        return ""
    lines = []
    for row in table_data:
        cells = [str(c).strip() if c else "" for c in row]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_pdf(file_path: str) -> str:
    doc = fitz.open(file_path)
    text_parts = []
    for page_num, page in enumerate(doc, 1):
        # Extract tables first (structured data is more reliable than get_text for tables)
        page_tables = page.find_tables()
        table_rects = []
        if page_tables and page_tables.tables:
            for table in page_tables.tables:
                try:
                    data = table.extract()
                    if data:
                        table_rects.append(table.bbox)
                        text_parts.append(
                            f"\n[Tabela - Pagina {page_num}]\n{_format_table(data)}\n"
                        )
                except Exception:
                    pass

        # Extract non-table text from the page
        text = page.get_text()
        if text.strip():
            text_parts.append(f"--- Pagina {page_num} ---\n{text}")

    doc.close()
    return "\n".join(text_parts)


def extract_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        text_parts.append("\n[Tabela]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text_parts.append(" | ".join(cells))
        text_parts.append("")

    return "\n".join(text_parts)


def extract_xlsx(file_path: str) -> str:
    wb = load_workbook(file_path, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"\n--- Aba: {sheet_name} ---")

        for row in ws.iter_rows(values_only=True):
            values = [str(v) if v is not None else "" for v in row]
            if any(v.strip() for v in values):
                text_parts.append(" | ".join(values))

    wb.close()
    return "\n".join(text_parts)


_IMAGE_TYPES = {"png", "jpg", "jpeg", "webp"}

# Abaixo deste total de caracteres um documento (nao-imagem) e tratado como
# "baixo rendimento" — tipicamente um PDF escaneado (so imagem) ou arquivo vazio.
# Um documento de engenharia real tem milhares de caracteres.
_MIN_CHARS_DOCUMENTO = 120


def aviso_extracao(file_type: str, texto: str | None) -> str | None:
    """Retorna um aviso quando o documento ficou sem texto suficiente — ou None.

    Rede contra a falha silenciosa: imagem/PDF escaneado entram como texto vazio
    e a analise rodaria no vazio sem o engenheiro saber. Derivado apenas do texto
    ja no banco — apos o OCR preencher `texto_extraido`, o aviso some sozinho
    (por isso o teste de rendimento vem ANTES do caso imagem).
    """
    chars = len((texto or "").strip())
    if chars >= _MIN_CHARS_DOCUMENTO:
        return None
    if file_type in _IMAGE_TYPES:
        return (
            "Imagem sem texto legivel: o OCR ainda nao recuperou conteudo (pode "
            "estar em andamento) ou a imagem esta ilegivel. Aguarde ou envie uma "
            "imagem mais nitida / um arquivo pesquisavel."
        )
    return (
        f"Pouco texto extraido ({chars} caracteres). Se este for um documento "
        "escaneado ou foto, o conteudo pode nao ter sido lido — aguarde o OCR ou "
        "envie um arquivo pesquisavel (PDF com texto, DOCX ou XLSX)."
    )


def extract_text(file_path: str, file_type: str) -> str:
    if file_type in _IMAGE_TYPES:
        # Imagens ficam anexadas ao caso. OCR pode ser conectado aqui no futuro.
        return ""

    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "xlsx": extract_xlsx,
    }

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Tipo de arquivo nao suportado: {file_type}")

    return extractor(file_path)
