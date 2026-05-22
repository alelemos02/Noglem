import io
import logging
from datetime import datetime

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
)
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.styles.protection import Protection
from openpyxl.utils import get_column_letter
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Coluna onde o fornecedor preenche a resposta (1-based: F = 6)
_CARTA_COL_RESPOSTA = 6
# Coluna oculta com ITEM_ID (G = 7)
_CARTA_COL_ITEM_ID = 7
# Senha de proteção da planilha (protege colunas A-E e G; deixa F editável)
_CARTA_SHEET_PASSWORD = "patec_pendencias"

logger = logging.getLogger(__name__)

# Shared constants for status/priority labels and colors
STATUS_LABELS = {
    "A": "APROVADO",
    "B": "APROVADO COM COMENTARIOS",
    "C": "REJEITADO",
    "D": "INFORMACAO AUSENTE",
    "E": "ITEM ADICIONAL DO FORNECEDOR",
}

STATUS_HEX_COLORS = {
    "A": "#c6f6d5",
    "B": "#fefcbf",
    "C": "#fed7d7",
    "D": "#e2e8f0",
    "E": "#bee3f8",
}

STATUS_ROW_HEX_COLORS = {
    "A": "#f0fff4",
    "B": "#fffff0",
    "C": "#fff5f5",
    "D": "#f7fafc",
    "E": "#ebf8ff",
}

MAX_EXPORT_ITEMS = 5000
DEFAULT_OBSERVATION_BY_STATUS = {
    "A": "Item aderente aos requisitos tecnicos da engenharia, sem desvios identificados.",
    "B": "Item parcialmente conforme; requer ajustes para atendimento completo.",
    "C": "Item nao conforme aos requisitos tecnicos e necessita revisao do fornecedor.",
    "D": "Informacao tecnica ausente na documentacao do fornecedor.",
    "E": "Item adicional apresentado pelo fornecedor e pendente de avaliacao da engenharia.",
}

_REPORT_TEXTS = {
    "pt": {
        "title": "PARECER TECNICO DE ENGENHARIA",
        "number": "Numero",
        "number_full": "Numero do Parecer",
        "date": "Data",
        "project": "Projeto",
        "revision": "Revisao",
        "supplier": "Fornecedor",
        "opinion": "Parecer",
        "general_opinion": "Parecer Geral",
        "pending": "Pendente",
        "status": "Status",
        "information": "Informacoes do Parecer",
        "executive_summary": "RESUMO EXECUTIVO",
        "executive_summary_title": "Resumo Executivo",
        "total": "Total",
        "total_items": "Total de Itens",
        "approved": "Aprovados",
        "approved_comments": "Aprov. c/ Com.",
        "approved_comments_full": "Aprovados com Comentarios",
        "rejected": "Rejeitados",
        "missing_info": "Info Ausente",
        "missing_info_full": "Informacao Ausente",
        "additional": "Adicionais",
        "additional_full": "Itens Adicionais",
        "general_comment": "Comentario Geral",
        "items": "ITENS DO PARECER",
        "items_title": "Itens do Parecer",
        "engineering_request": "Solicitacao da Engenharia",
        "supplier_proposal": "Proposto pelo Fornecedor",
        "observation": "Observacao",
        "supplier_comment": "Comentario Fornecedor",
        "required_value": "Valor requerido",
        "engineering_ref": "Ref. engenharia",
        "informed_value": "Valor informado",
        "supplier_ref": "Ref. fornecedor",
        "not_informed": "Nao informado",
        "conclusion": "CONCLUSAO",
        "conclusion_title": "Conclusao",
        "recommendations": "RECOMENDACOES",
        "recommendations_title": "Recomendacoes",
        "sheet_summary": "Resumo",
        "sheet_items": "Itens",
        "sheet_recommendations": "Recomendacoes",
    },
    "es": {
        "title": "DICTAMEN TECNICO DE INGENIERIA",
        "number": "Numero",
        "number_full": "Numero del Dictamen",
        "date": "Fecha",
        "project": "Proyecto",
        "revision": "Revision",
        "supplier": "Proveedor",
        "opinion": "Dictamen",
        "general_opinion": "Dictamen General",
        "pending": "Pendiente",
        "status": "Estado",
        "information": "Informacion del Dictamen",
        "executive_summary": "RESUMEN EJECUTIVO",
        "executive_summary_title": "Resumen Ejecutivo",
        "total": "Total",
        "total_items": "Total de Items",
        "approved": "Aprobados",
        "approved_comments": "Aprob. c/ Com.",
        "approved_comments_full": "Aprobados con Comentarios",
        "rejected": "Rechazados",
        "missing_info": "Info Ausente",
        "missing_info_full": "Informacion Ausente",
        "additional": "Adicionales",
        "additional_full": "Items Adicionales",
        "general_comment": "Comentario General",
        "items": "ITEMS DEL DICTAMEN",
        "items_title": "Items del Dictamen",
        "engineering_request": "Solicitud de Ingenieria",
        "supplier_proposal": "Propuesto por el Proveedor",
        "observation": "Observacion",
        "supplier_comment": "Comentario Proveedor",
        "required_value": "Valor requerido",
        "engineering_ref": "Ref. ingenieria",
        "informed_value": "Valor informado",
        "supplier_ref": "Ref. proveedor",
        "not_informed": "No informado",
        "conclusion": "CONCLUSION",
        "conclusion_title": "Conclusion",
        "recommendations": "RECOMENDACIONES",
        "recommendations_title": "Recomendaciones",
        "sheet_summary": "Resumen",
        "sheet_items": "Items",
        "sheet_recommendations": "Recomendaciones",
    },
    "en": {
        "title": "ENGINEERING TECHNICAL OPINION",
        "number": "Number",
        "number_full": "Opinion Number",
        "date": "Date",
        "project": "Project",
        "revision": "Revision",
        "supplier": "Supplier",
        "opinion": "Opinion",
        "general_opinion": "Overall Opinion",
        "pending": "Pending",
        "status": "Status",
        "information": "Opinion Information",
        "executive_summary": "EXECUTIVE SUMMARY",
        "executive_summary_title": "Executive Summary",
        "total": "Total",
        "total_items": "Total Items",
        "approved": "Approved",
        "approved_comments": "Appr. w/ Com.",
        "approved_comments_full": "Approved with Comments",
        "rejected": "Rejected",
        "missing_info": "Missing Info",
        "missing_info_full": "Missing Information",
        "additional": "Additional",
        "additional_full": "Additional Items",
        "general_comment": "General Comment",
        "items": "OPINION ITEMS",
        "items_title": "Opinion Items",
        "engineering_request": "Engineering Request",
        "supplier_proposal": "Supplier Proposal",
        "observation": "Observation",
        "supplier_comment": "Supplier Comment",
        "required_value": "Required value",
        "engineering_ref": "Engineering ref.",
        "informed_value": "Informed value",
        "supplier_ref": "Supplier ref.",
        "not_informed": "Not informed",
        "conclusion": "CONCLUSION",
        "conclusion_title": "Conclusion",
        "recommendations": "RECOMMENDATIONS",
        "recommendations_title": "Recommendations",
        "sheet_summary": "Summary",
        "sheet_items": "Items",
        "sheet_recommendations": "Recommendations",
    },
}

_STATUS_LABELS_BY_LANGUAGE = {
    "pt": STATUS_LABELS,
    "es": {
        "A": "APROBADO",
        "B": "APROBADO CON COMENTARIOS",
        "C": "RECHAZADO",
        "D": "INFORMACION AUSENTE",
        "E": "ITEM ADICIONAL DEL PROVEEDOR",
    },
    "en": {
        "A": "APPROVED",
        "B": "APPROVED WITH COMMENTS",
        "C": "REJECTED",
        "D": "MISSING INFORMATION",
        "E": "ADDITIONAL SUPPLIER ITEM",
    },
}

_DEFAULT_OBSERVATIONS_BY_LANGUAGE = {
    "pt": DEFAULT_OBSERVATION_BY_STATUS,
    "es": {
        "A": "Item conforme con los requisitos tecnicos de ingenieria, sin desvios identificados.",
        "B": "Item parcialmente conforme; requiere ajustes para cumplimiento completo.",
        "C": "Item no conforme con los requisitos tecnicos y requiere revision del proveedor.",
        "D": "Informacion tecnica ausente en la documentacion del proveedor.",
        "E": "Item adicional presentado por el proveedor y pendiente de evaluacion de ingenieria.",
    },
    "en": {
        "A": "Item meets the engineering technical requirements with no deviations identified.",
        "B": "Item is partially compliant and requires adjustments for full compliance.",
        "C": "Item is non-compliant with the technical requirements and requires supplier revision.",
        "D": "Technical information is missing from the supplier documentation.",
        "E": "Additional supplier item pending engineering evaluation.",
    },
}

_GENERAL_OPINION_LABELS = {
    "pt": {
        "APROVADO": "APROVADO",
        "APROVADO COM COMENTARIOS": "APROVADO COM COMENTARIOS",
        "REJEITADO": "REJEITADO",
    },
    "es": {
        "APROVADO": "APROBADO",
        "APROVADO COM COMENTARIOS": "APROBADO CON COMENTARIOS",
        "REJEITADO": "RECHAZADO",
    },
    "en": {
        "APROVADO": "APPROVED",
        "APROVADO COM COMENTARIOS": "APPROVED WITH COMMENTS",
        "REJEITADO": "REJECTED",
    },
}

_PROCESSING_STATUS_LABELS = {
    "pt": {
        "pendente": "Pendente",
        "processando": "Processando",
        "concluido": "Concluido",
        "erro": "Erro",
    },
    "es": {
        "pendente": "Pendiente",
        "processando": "Procesando",
        "concluido": "Concluido",
        "erro": "Error",
    },
    "en": {
        "pendente": "Pending",
        "processando": "Processing",
        "concluido": "Completed",
        "erro": "Error",
    },
}


def _get_report_language(parecer) -> str:
    idioma = getattr(parecer, "idioma_relatorio", "pt")
    return idioma if idioma in _REPORT_TEXTS else "pt"


def _get_report_texts(idioma_relatorio: str) -> dict[str, str]:
    return _REPORT_TEXTS.get(idioma_relatorio, _REPORT_TEXTS["pt"])


def _status_label(status: str, idioma_relatorio: str) -> str:
    return _STATUS_LABELS_BY_LANGUAGE.get(idioma_relatorio, STATUS_LABELS).get(status, status)


def _general_opinion_label(parecer_geral: str | None, idioma_relatorio: str) -> str:
    texts = _get_report_texts(idioma_relatorio)
    if not parecer_geral:
        return texts["pending"]
    return _GENERAL_OPINION_LABELS.get(idioma_relatorio, _GENERAL_OPINION_LABELS["pt"]).get(
        parecer_geral,
        parecer_geral,
    )


def _processing_status_label(status_processamento: str, idioma_relatorio: str) -> str:
    return _PROCESSING_STATUS_LABELS.get(idioma_relatorio, _PROCESSING_STATUS_LABELS["pt"]).get(
        status_processamento,
        status_processamento,
    )


def _build_solicitacao_engenharia(item, idioma_relatorio: str = "pt") -> str:
    texts = _get_report_texts(idioma_relatorio)
    parts = [item.descricao_requisito or ""]
    if item.valor_requerido:
        parts.append(f"{texts['required_value']}: {item.valor_requerido}")
    if item.referencia_engenharia:
        parts.append(f"{texts['engineering_ref']}: {item.referencia_engenharia}")
    return "\n".join(part for part in parts if part and part.strip()).strip() or "-"


def _build_proposto_fornecedor(item, idioma_relatorio: str = "pt") -> str:
    texts = _get_report_texts(idioma_relatorio)
    parts = []
    if item.valor_fornecedor:
        parts.append(f"{texts['informed_value']}: {item.valor_fornecedor}")
    if item.referencia_fornecedor:
        parts.append(f"{texts['supplier_ref']}: {item.referencia_fornecedor}")
    return "\n".join(part for part in parts if part and part.strip()).strip() or texts["not_informed"]


def _build_observacao(item, idioma_relatorio: str = "pt") -> str:
    justificativa = (item.justificativa_tecnica or "").strip()
    if justificativa:
        return justificativa

    defaults = _DEFAULT_OBSERVATIONS_BY_LANGUAGE.get(
        idioma_relatorio,
        DEFAULT_OBSERVATION_BY_STATUS,
    )
    return defaults.get(
        item.status,
        _DEFAULT_OBSERVATIONS_BY_LANGUAGE["pt"]["B"],
    )


def export_pdf(parecer, itens, recomendacoes) -> bytes:
    """Generate a PDF report for the parecer."""
    try:
        export_itens = itens[:MAX_EXPORT_ITEMS]
        idioma_relatorio = _get_report_language(parecer)
        texts = _get_report_texts(idioma_relatorio)
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            leftMargin=20 * mm, rightMargin=20 * mm,
            topMargin=25 * mm, bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name="SmallBody", parent=styles["Normal"], fontSize=8, leading=10))
        styles.add(ParagraphStyle(
            name="Header1", parent=styles["Heading1"], fontSize=14,
            spaceAfter=6, textColor=colors.HexColor("#1a365d"),
        ))
        styles.add(ParagraphStyle(
            name="Header2", parent=styles["Heading2"], fontSize=11,
            spaceAfter=4, textColor=colors.HexColor("#2d3748"),
        ))
        styles.add(ParagraphStyle(
            name="CenteredSmallBody", parent=styles["SmallBody"], fontSize=7, leading=9, alignment=1
        ))

        elements = []

        # Title
        elements.append(Paragraph(texts["title"], styles["Header1"]))
        elements.append(Spacer(1, 4 * mm))

        # Metadata table
        meta_data = [
            [f"{texts['number']}:", parecer.numero_parecer, f"{texts['date']}:", datetime.now().strftime("%d/%m/%Y")],
            [f"{texts['project']}:", parecer.projeto, f"{texts['revision']}:", parecer.revisao],
            [f"{texts['supplier']}:", parecer.fornecedor, f"{texts['opinion']}:", _general_opinion_label(parecer.parecer_geral, idioma_relatorio)],
        ]
        meta_table = Table(meta_data, colWidths=[25 * mm, 60 * mm, 22 * mm, 60 * mm])
        meta_table.setStyle(TableStyle([
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f7fafc")),
            ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#f7fafc")),
        ]))
        elements.append(meta_table)
        elements.append(Spacer(1, 6 * mm))

        # Executive Summary
        elements.append(Paragraph(texts["executive_summary"], styles["Header2"]))
        summary_data = [
            [
                texts["total"],
                texts["approved"],
                texts["approved_comments"],
                texts["rejected"],
                texts["missing_info"],
                texts["additional"],
            ],
            [
                str(parecer.total_itens), str(parecer.total_aprovados),
                str(parecer.total_aprovados_comentarios), str(parecer.total_rejeitados),
                str(parecer.total_info_ausente), str(parecer.total_itens_adicionais),
            ],
        ]
        summary_table = Table(summary_data, colWidths=[28 * mm] * 6)

        summary_col_colors = [
            (1, colors.HexColor(STATUS_HEX_COLORS["A"])),
            (2, colors.HexColor(STATUS_HEX_COLORS["B"])),
            (3, colors.HexColor(STATUS_HEX_COLORS["C"])),
            (4, colors.HexColor(STATUS_HEX_COLORS["D"])),
            (5, colors.HexColor(STATUS_HEX_COLORS["E"])),
        ]
        style_cmds = [
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#edf2f7")),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
        ]
        for col, color in summary_col_colors:
            style_cmds.append(("BACKGROUND", (col, 1), (col, 1), color))

        summary_table.setStyle(TableStyle(style_cmds))
        elements.append(summary_table)
        elements.append(Spacer(1, 4 * mm))

        if parecer.comentario_geral:
            elements.append(Paragraph(
                f"<b>{texts['general_comment']}:</b> {parecer.comentario_geral}", styles["SmallBody"]
            ))
            elements.append(Spacer(1, 4 * mm))

        # Items table
        if export_itens:
            elements.append(Paragraph(texts["items"], styles["Header2"]))

            item_header = [
                texts["engineering_request"],
                texts["supplier_proposal"],
                texts["status"],
                texts["observation"],
            ]
            col_widths = [50 * mm, 42 * mm, 22 * mm, 56 * mm]

            rows = [item_header]
            for item in export_itens:
                rows.append([
                    Paragraph(_build_solicitacao_engenharia(item, idioma_relatorio)[:800], styles["SmallBody"]),
                    Paragraph(_build_proposto_fornecedor(item, idioma_relatorio)[:800], styles["SmallBody"]),
                    Paragraph(_status_label(item.status, idioma_relatorio), styles["CenteredSmallBody"]),
                    Paragraph(_build_observacao(item, idioma_relatorio)[:800], styles["SmallBody"]),
                ])

            item_table = Table(rows, colWidths=col_widths, repeatRows=1)

            row_styles = [
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2d3748")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("ALIGN", (2, 0), (2, -1), "CENTER"),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
            ]

            for i, item in enumerate(export_itens, 1):
                hex_color = STATUS_ROW_HEX_COLORS.get(item.status)
                if hex_color:
                    row_styles.append(("BACKGROUND", (0, i), (-1, i), colors.HexColor(hex_color)))

            item_table.setStyle(TableStyle(row_styles))
            elements.append(item_table)
            elements.append(Spacer(1, 6 * mm))

        # Conclusion
        if parecer.conclusao:
            elements.append(Paragraph(texts["conclusion"], styles["Header2"]))
            elements.append(Paragraph(parecer.conclusao, styles["Normal"]))
            elements.append(Spacer(1, 4 * mm))

        # Recommendations
        if recomendacoes:
            elements.append(Paragraph(texts["recommendations"], styles["Header2"]))
            for rec in recomendacoes:
                elements.append(Paragraph(f"&bull; {rec.texto}", styles["Normal"]))
            elements.append(Spacer(1, 4 * mm))

        doc.build(elements)
        return buffer.getvalue()

    except Exception:
        logger.exception("Erro ao gerar PDF para parecer %s", parecer.numero_parecer)
        raise


def export_xlsx(parecer, itens, recomendacoes) -> bytes:
    """Generate an XLSX report for the parecer."""
    try:
        export_itens = itens[:MAX_EXPORT_ITEMS]
        idioma_relatorio = _get_report_language(parecer)
        texts = _get_report_texts(idioma_relatorio)
        wb = Workbook()

        header_font = Font(bold=True, color="FFFFFF", size=10)
        header_fill = PatternFill(start_color="2D3748", end_color="2D3748", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )

        status_fills = {
            code: PatternFill(
                start_color=hex_color.lstrip("#"),
                end_color=hex_color.lstrip("#"),
                fill_type="solid",
            )
            for code, hex_color in STATUS_HEX_COLORS.items()
        }

        # Sheet 1: Resumo
        ws = wb.active
        ws.title = texts["sheet_summary"]
        ws.column_dimensions["A"].width = 25
        ws.column_dimensions["B"].width = 50

        info = [
            (texts["title"], ""),
            ("", ""),
            (texts["number_full"], parecer.numero_parecer),
            (texts["project"], parecer.projeto),
            (texts["supplier"], parecer.fornecedor),
            (texts["revision"], parecer.revisao),
            (texts["date"], datetime.now().strftime("%d/%m/%Y")),
            (texts["general_opinion"], _general_opinion_label(parecer.parecer_geral, idioma_relatorio)),
            ("", ""),
            (texts["executive_summary"], ""),
            (texts["total_items"], parecer.total_itens),
            (f"{texts['approved']} (A)", parecer.total_aprovados),
            (f"{texts['approved_comments_full']} (B)", parecer.total_aprovados_comentarios),
            (f"{texts['rejected']} (C)", parecer.total_rejeitados),
            (f"{texts['missing_info_full']} (D)", parecer.total_info_ausente),
            (f"{texts['additional_full']} (E)", parecer.total_itens_adicionais),
        ]
        for row_idx, (label, value) in enumerate(info, 1):
            ws.cell(row=row_idx, column=1, value=label).font = Font(bold=True) if label else Font()
            ws.cell(row=row_idx, column=2, value=value)

        ws.cell(row=1, column=1).font = Font(bold=True, size=14)

        if parecer.comentario_geral:
            ws.append([""])
            ws.append([texts["general_comment"], parecer.comentario_geral])

        if parecer.conclusao:
            ws.append([""])
            ws.append([texts["conclusion_title"], parecer.conclusao])

        # Sheet 2: Itens
        ws2 = wb.create_sheet(texts["sheet_items"])
        headers = [
            texts["engineering_request"],
            texts["supplier_proposal"],
            texts["status"],
            texts["observation"],
            texts["supplier_comment"],
        ]
        supplier_comment_col = len(headers)  # last column index (1-based)
        supplier_comment_fill = PatternFill(start_color="FFF9C4", end_color="FFF9C4", fill_type="solid")
        for col, h in enumerate(headers, 1):
            cell = ws2.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill if col != supplier_comment_col else PatternFill(
                start_color="F9A825", end_color="F9A825", fill_type="solid"
            )
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center", wrap_text=True)

        col_widths = [60, 55, 24, 75, 55]
        for i, w in enumerate(col_widths, 1):
            ws2.column_dimensions[get_column_letter(i)].width = w

        for item in export_itens:
            row = [
                _build_solicitacao_engenharia(item, idioma_relatorio),
                _build_proposto_fornecedor(item, idioma_relatorio),
                _status_label(item.status, idioma_relatorio),
                _build_observacao(item, idioma_relatorio),
                "",
            ]
            ws2.append(row)
            row_idx = ws2.max_row
            fill = status_fills.get(item.status)
            for col in range(1, len(headers) + 1):
                cell = ws2.cell(row=row_idx, column=col)
                if col == supplier_comment_col:
                    cell.fill = supplier_comment_fill
                elif fill:
                    cell.fill = fill
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="top")

        # Sheet 3: Recomendacoes
        ws3 = wb.create_sheet(texts["sheet_recommendations"])
        for col, h in enumerate(["#", texts["recommendations_title"]], 1):
            cell = ws3.cell(row=1, column=col, value=h)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = thin_border
        ws3.column_dimensions["A"].width = 5
        ws3.column_dimensions["B"].width = 80

        for rec in recomendacoes:
            ws3.append([rec.ordem, rec.texto])

        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()

    except Exception:
        logger.exception("Erro ao gerar XLSX para parecer %s", parecer.numero_parecer)
        raise


def export_carta_pendencias(parecer, itens_pendentes) -> bytes:
    """
    Gera a carta de pendências em XLSX para os itens em PENDENTE_FORNECEDOR.

    Estrutura:
      - Linha 1: instrução ao fornecedor
      - Linha 2: cabeçalhos (A–G)
      - Linhas 3+: um item por linha
    Colunas A–E e G são protegidas (somente leitura).
    Coluna F (Resposta do Fornecedor) é editável.
    Coluna G (ITEM_ID) é oculta para guiar o reimport determinístico.
    """
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Pendencias"

        header_font = Font(bold=True, color="FFFFFF", size=10)
        header_fill = PatternFill(start_color="2D3748", end_color="2D3748", fill_type="solid")
        instrucao_fill = PatternFill(start_color="FFF3CD", end_color="FFF3CD", fill_type="solid")
        resposta_fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
        thin_border = Border(
            left=Side(style="thin"), right=Side(style="thin"),
            top=Side(style="thin"), bottom=Side(style="thin"),
        )
        locked_protection = Protection(locked=True)
        unlocked_protection = Protection(locked=False)

        # ── Linha 1: instrução ──────────────────────────────────────────
        instrucao = (
            "INSTRUCAO AO FORNECEDOR: Preencha APENAS a coluna F (Resposta do Fornecedor). "
            "Nao altere nenhuma outra celula. Nao remova linhas nem colunas. "
            "Retorne este arquivo preenchido para avaliacao."
        )
        ws.merge_cells("A1:F1")
        cell = ws["A1"]
        cell.value = instrucao
        cell.font = Font(bold=True, size=9, color="7B4F00")
        cell.fill = instrucao_fill
        cell.alignment = Alignment(wrap_text=True, vertical="center")
        cell.protection = locked_protection
        ws.row_dimensions[1].height = 36

        # ── Linha 2: cabeçalhos ─────────────────────────────────────────
        cabecalhos = [
            "Item",
            "Disciplina/Categoria",
            "Prioridade",
            "Requisito da Engenharia",
            "Pendencia / Acao Requerida",
            "Resposta do Fornecedor",
            "ITEM_ID",
        ]
        for col, titulo in enumerate(cabecalhos, 1):
            cell = ws.cell(row=2, column=col, value=titulo)
            cell.font = header_font
            cell.fill = header_fill if col != _CARTA_COL_RESPOSTA else PatternFill(
                start_color="1B5E20", end_color="1B5E20", fill_type="solid"
            )
            cell.border = thin_border
            cell.alignment = Alignment(horizontal="center", wrap_text=True)
            cell.protection = locked_protection

        # Larguras das colunas
        col_widths = [8, 22, 14, 55, 55, 55, 15]
        for i, w in enumerate(col_widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w

        # Oculta coluna ITEM_ID
        ws.column_dimensions[get_column_letter(_CARTA_COL_ITEM_ID)].hidden = True

        # ── Linhas de dados ─────────────────────────────────────────────
        prioridade_colors = {"ALTA": "FFCCCC", "MEDIA": "FFF3CC", "BAIXA": "E8F5E9"}

        for item in itens_pendentes:
            # Busca acao_requerida da última rodada (campo do item, fallback para o campo direto)
            pendencia = (item.acao_requerida or "").strip() or "Ver justificativa tecnica."

            row_values = [
                item.numero,
                item.categoria or "",
                item.prioridade or "",
                (item.descricao_requisito or "").strip(),
                pendencia,
                "",                      # Resposta — editável
                str(item.id),            # ITEM_ID — oculto
            ]

            ws.append(row_values)
            row_idx = ws.max_row

            prio_color = prioridade_colors.get(item.prioridade or "", "FFFFFF")

            for col in range(1, _CARTA_COL_ITEM_ID + 1):
                cell = ws.cell(row=row_idx, column=col)
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical="top")

                if col == _CARTA_COL_RESPOSTA:
                    cell.fill = resposta_fill
                    cell.protection = unlocked_protection
                elif col == 3:  # prioridade
                    cell.fill = PatternFill(
                        start_color=prio_color, end_color=prio_color, fill_type="solid"
                    )
                    cell.protection = locked_protection
                else:
                    cell.protection = locked_protection

        # ── Metadados do parecer numa aba separada ──────────────────────
        ws_meta = wb.create_sheet("Info")
        ws_meta["A1"] = "Parecer"
        ws_meta["B1"] = parecer.numero_parecer
        ws_meta["A2"] = "Projeto"
        ws_meta["B2"] = parecer.projeto
        ws_meta["A3"] = "Fornecedor"
        ws_meta["B3"] = parecer.fornecedor
        ws_meta["A4"] = "Rodada"
        ws_meta["B4"] = getattr(parecer, "rodada_atual", 1)
        ws_meta["A5"] = "Data emissao"
        ws_meta["B5"] = datetime.now().strftime("%d/%m/%Y")
        ws_meta["A6"] = "Total pendentes"
        ws_meta["B6"] = len(itens_pendentes)
        ws_meta.column_dimensions["A"].width = 20
        ws_meta.column_dimensions["B"].width = 40

        # ── Protege a aba Pendencias (somente col F editável) ───────────
        ws.protection.sheet = True
        ws.protection.password = _CARTA_SHEET_PASSWORD
        ws.protection.enable()

        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()

    except Exception:
        logger.exception("Erro ao gerar carta de pendencias para parecer %s", parecer.numero_parecer)
        raise


def export_docx(parecer, itens, recomendacoes) -> bytes:
    """Generate a DOCX report for the parecer."""
    try:
        return _build_docx(parecer, itens, recomendacoes)
    except Exception:
        logger.exception("Erro ao gerar DOCX para parecer %s", parecer.numero_parecer)
        raise


def _build_docx(parecer, itens, recomendacoes) -> bytes:
    export_itens = itens[:MAX_EXPORT_ITEMS]
    idioma_relatorio = _get_report_language(parecer)
    texts = _get_report_texts(idioma_relatorio)
    doc = DocxDocument()

    # Style adjustments
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    style.font.name = "Calibri"

    # Title
    title = doc.add_heading(texts["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_heading(texts["information"], level=1)
    meta_table = doc.add_table(rows=4, cols=4)
    meta_table.style = "Table Grid"
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    meta_data = [
        (f"{texts['number']}:", parecer.numero_parecer, f"{texts['date']}:", datetime.now().strftime("%d/%m/%Y")),
        (f"{texts['project']}:", parecer.projeto, f"{texts['revision']}:", parecer.revisao),
        (f"{texts['supplier']}:", parecer.fornecedor, f"{texts['opinion']}:", _general_opinion_label(parecer.parecer_geral, idioma_relatorio)),
        (f"{texts['total_items']}:", str(parecer.total_itens), f"{texts['status']}:", _processing_status_label(parecer.status_processamento, idioma_relatorio)),
    ]
    for i, (l1, v1, l2, v2) in enumerate(meta_data):
        row = meta_table.rows[i]
        row.cells[0].text = l1
        row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
        row.cells[1].text = v1
        row.cells[2].text = l2
        row.cells[3].text = v2
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)

    # Bold labels
    for i in range(4):
        for col in [0, 2]:
            for run in meta_table.rows[i].cells[col].paragraphs[0].runs:
                run.bold = True

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading(texts["executive_summary_title"], level=1)
    summary_table = doc.add_table(rows=2, cols=6)
    summary_table.style = "Table Grid"
    summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        texts["total"],
        texts["approved"],
        texts["approved_comments"],
        texts["rejected"],
        texts["missing_info"],
        texts["additional"],
    ]
    values = [
        parecer.total_itens, parecer.total_aprovados, parecer.total_aprovados_comentarios,
        parecer.total_rejeitados, parecer.total_info_ausente, parecer.total_itens_adicionais,
    ]
    for i, h in enumerate(headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, v in enumerate(values):
        cell = summary_table.rows[1].cells[i]
        cell.text = str(v)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if parecer.comentario_geral:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(f"{texts['general_comment']}: ").bold = True
        p.add_run(parecer.comentario_geral)

    doc.add_paragraph()

    # Items
    if export_itens:
        doc.add_heading(texts["items_title"], level=1)

        items_table = doc.add_table(rows=1, cols=4)
        items_table.style = "Table Grid"
        items_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        headers = [
            texts["engineering_request"],
            texts["supplier_proposal"],
            texts["status"],
            texts["observation"],
        ]
        for idx, title in enumerate(headers):
            cell = items_table.rows[0].cells[idx]
            cell.text = title
            for run in cell.paragraphs[0].runs:
                run.bold = True

        for item in export_itens:
            row = items_table.add_row().cells
            row[0].text = _build_solicitacao_engenharia(item, idioma_relatorio)
            row[1].text = _build_proposto_fornecedor(item, idioma_relatorio)
            row[2].text = _status_label(item.status, idioma_relatorio)
            row[3].text = _build_observacao(item, idioma_relatorio)

    # Conclusion
    if parecer.conclusao:
        doc.add_heading(texts["conclusion_title"], level=1)
        doc.add_paragraph(parecer.conclusao)

    # Recommendations
    if recomendacoes:
        doc.add_heading(texts["recommendations_title"], level=1)
        for rec in recomendacoes:
            doc.add_paragraph(rec.texto, style="List Bullet")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
