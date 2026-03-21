"""Text extraction for Conhecimento RAG documents.

Supports PDF (with OCR fallback for scanned documents), DOCX, and XLSX.
Based on PATEC's text_extractor with OCR fallback from the old RAG service.
"""

import logging

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

# Minimum characters from standard extraction before triggering OCR
OCR_THRESHOLD = 50


def _format_table(table_data: list[list]) -> str:
    """Format extracted table data as pipe-separated text with headers."""
    if not table_data:
        return ""
    lines = []
    for row in table_data:
        cells = [str(c).strip() if c else "" for c in row]
        lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_pdf(file_path: str) -> tuple[str, bool]:
    """Extract text from a PDF file.

    Uses PyMuPDF for standard extraction. Falls back to OCR (Tesseract)
    if the extracted text is below the threshold (scanned documents).

    Returns:
        Tuple of (extracted_text, has_ocr).
    """
    doc = fitz.open(file_path)
    text_parts = []
    total_length = 0

    for page_num, page in enumerate(doc, 1):
        # Extract tables first (structured data is more reliable)
        page_tables = page.find_tables()
        if page_tables and page_tables.tables:
            for table in page_tables.tables:
                try:
                    data = table.extract()
                    if data:
                        text_parts.append(
                            f"\n[Tabela - Pagina {page_num}]\n{_format_table(data)}\n"
                        )
                except Exception:
                    pass

        # Extract non-table text from the page
        text = page.get_text()
        if text.strip():
            text_parts.append(f"--- Pagina {page_num} ---\n{text}")
            total_length += len(text)

    doc.close()

    # Check if OCR is needed (scanned PDF with insufficient text)
    if total_length < OCR_THRESHOLD:
        logger.info("PDF text below threshold (%d chars), attempting OCR fallback", total_length)
        ocr_text = _extract_pdf_ocr(file_path)
        if ocr_text:
            return ocr_text, True

    return "\n".join(text_parts), False


def _extract_pdf_ocr(file_path: str) -> str | None:
    """Extract text from a scanned PDF using Tesseract OCR.

    Returns extracted text with page markers, or None if OCR fails.
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError:
        logger.warning("OCR dependencies (pdf2image, pytesseract) not installed, skipping OCR")
        return None

    try:
        images = convert_from_path(file_path)
        text_parts = []

        for i, image in enumerate(images):
            text = pytesseract.image_to_string(image)
            if text.strip():
                text_parts.append(f"--- Pagina {i + 1} ---\n{text}")

        if text_parts:
            return "\n".join(text_parts)

    except Exception as e:
        logger.error("OCR extraction failed: %s", e)

    return None


def extract_docx(file_path: str) -> tuple[str, bool]:
    """Extract text from a DOCX file.

    Returns:
        Tuple of (extracted_text, has_ocr=False).
    """
    doc = DocxDocument(file_path)
    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        text_parts.append("\n[Tabela]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            text_parts.append(" | ".join(cells))
        text_parts.append("")

    return "\n".join(text_parts), False


def extract_xlsx(file_path: str) -> tuple[str, bool]:
    """Extract text from an XLSX file.

    Returns:
        Tuple of (extracted_text, has_ocr=False).
    """
    wb = load_workbook(file_path, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text_parts.append(f"\n--- Aba: {sheet_name} ---")

        for row in ws.iter_rows(values_only=True):
            values = [str(v) if v is not None else "" for v in row]
            if any(v.strip() for v in values):
                text_parts.append(" | ".join(values))

    wb.close()
    return "\n".join(text_parts), False


def extract_text(file_path: str, file_type: str) -> tuple[str, bool]:
    """Extract text from a file based on its type.

    Args:
        file_path: Path to the file.
        file_type: One of "pdf", "docx", "xlsx".

    Returns:
        Tuple of (extracted_text, has_ocr).
    """
    extractors = {
        "pdf": extract_pdf,
        "docx": extract_docx,
        "xlsx": extract_xlsx,
    }

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Tipo de arquivo nao suportado: {file_type}")

    return extractor(file_path)
